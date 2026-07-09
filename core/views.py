from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.views.generic import View
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, JsonResponse
from django.conf import settings
from translation.services import SUPPORTED_LANGUAGES


@login_required(login_url='accounts:login')
def home(request):
    if request.user.is_authenticated:
        return redirect('dashboard:dashboard')
    return render(request, 'home.html')


class LiveTranslateView(LoginRequiredMixin, View):
    template_name = 'live_translate.html'

    def get(self, request, *args, **kwargs):
        return render(request, self.template_name, {
            'languages': SUPPORTED_LANGUAGES,
        })


class DownloadTXTView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        import json
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        original = data.get('original', '')
        translated = data.get('translated', '')
        lang = data.get('language', '')

        lines = []
        lines.append("=== AI Speech Translation ===")
        lines.append(f"Date: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"Language: {SUPPORTED_LANGUAGES.get(lang, lang)}")
        lines.append("")
        lines.append("--- Original Text ---")
        lines.append(original)
        lines.append("")
        lines.append("--- Translated Text ---")
        lines.append(translated)

        content = '\n'.join(lines)
        resp = HttpResponse(content, content_type='text/plain; charset=utf-8')
        resp['Content-Disposition'] = 'attachment; filename="transcript.txt"'
        return resp


class DownloadSRTView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        import json
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        original = data.get('original', '')
        translated = data.get('translated', '')

        lines = []
        lines.append("1")
        lines.append("00:00:00,000 --> 00:01:00,000")
        lines.append(f"Original: {original}")
        lines.append(f"Translated: {translated}")
        lines.append("")

        resp = HttpResponse('\n'.join(lines), content_type='text/plain; charset=utf-8')
        resp['Content-Disposition'] = 'attachment; filename="subtitles.srt"'
        return resp


class DownloadJSONView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        import json
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        from datetime import datetime
        output = {
            'app': 'AI Speech Translation',
            'date': datetime.now().isoformat(),
            'language': SUPPORTED_LANGUAGES.get(data.get('language', ''), data.get('language', '')),
            'original_text': data.get('original', ''),
            'translated_text': data.get('translated', ''),
            'word_count_original': len(data.get('original', '').split()),
            'word_count_translated': len(data.get('translated', '').split()),
        }

        resp = HttpResponse(
            json.dumps(output, indent=2, ensure_ascii=False),
            content_type='application/json; charset=utf-8'
        )
        resp['Content-Disposition'] = 'attachment; filename="transcript.json"'
        return resp


class DownloadPDFView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        import json
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        original = data.get('original', '')
        translated = data.get('translated', '')
        lang = data.get('language', '')
        import datetime

        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        from io import BytesIO
        buffer = BytesIO()

        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            topMargin=20*mm, bottomMargin=20*mm,
            leftMargin=20*mm, rightMargin=20*mm
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Title'],
            fontSize=20, spaceAfter=20,
            textColor=HexColor('#667eea')
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'],
            fontSize=14, spaceAfter=10,
            textColor=HexColor('#764ba2')
        )
        body_style = ParagraphStyle(
            'CustomBody', parent=styles['Normal'],
            fontSize=11, spaceAfter=6,
            leading=16
        )
        meta_style = ParagraphStyle(
            'Meta', parent=styles['Normal'],
            fontSize=9, textColor=HexColor('#666666'),
            spaceAfter=4
        )

        elements = []
        elements.append(Paragraph("AI Speech Translation", title_style))
        elements.append(Spacer(1, 6*mm))
        elements.append(Paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", meta_style))
        elements.append(Paragraph(f"Language: {SUPPORTED_LANGUAGES.get(lang, lang)}", meta_style))
        elements.append(Spacer(1, 6*mm))
        elements.append(Paragraph("Original Text", heading_style))
        elements.append(Paragraph(original.replace('\n', '<br/>'), body_style))
        elements.append(Spacer(1, 4*mm))
        elements.append(Paragraph("Translated Text", heading_style))
        elements.append(Paragraph(translated.replace('\n', '<br/>'), body_style))

        doc.build(elements)
        pdf = buffer.getvalue()
        buffer.close()

        resp = HttpResponse(pdf, content_type='application/pdf')
        resp['Content-Disposition'] = 'attachment; filename="transcript.pdf"'
        return resp


class DownloadDOCXView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        import json
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        original = data.get('original', '')
        translated = data.get('translated', '')
        lang = data.get('language', '')
        import datetime

        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('AI Speech Translation')
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)

        doc.add_paragraph(f'Date: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Language: {SUPPORTED_LANGUAGES.get(lang, lang)}')
        doc.add_paragraph('')

        h = doc.add_heading('Original Text', level=2)
        doc.add_paragraph(original)

        h = doc.add_heading('Translated Text', level=2)
        doc.add_paragraph(translated)

        buffer = __import__('io').BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()
        buffer.close()

        resp = HttpResponse(content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = 'attachment; filename="transcript.docx"'
        return resp
